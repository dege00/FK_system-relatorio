"""
Módulo responsável pela manipulação do template Word.
Gerencia substituição de placeholders, inserção de imagens e duplicação de páginas.
"""

import os
import shutil
import logging
from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.shared import Emu, Inches
from PIL import Image as PILImage

import modules.photo_handler as ph

logger = logging.getLogger(__name__)

# Adiciona namespace VML necessário para manipulação de imagens
nsmap['v'] = 'urn:schemas-microsoft-com:vml'
nsmap['wp'] = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
nsmap['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'
nsmap['pic'] = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
nsmap['r'] = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# Namespaces completos para uso direto
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
NS_WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
NS_A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
NS_PIC = 'http://schemas.openxmlformats.org/drawingml/2006/picture'

# Constantes
PHOTO_ROWS_PER_PAGE = 2
PHOTOS_PER_ROW = 2
PHOTOS_PER_PAGE = PHOTO_ROWS_PER_PAGE * PHOTOS_PER_ROW


class TemplateHandler:
    """
    Gerencia a abertura, preenchimento e salvamento do template Word.
    """

    def __init__(self, caminho_template: str):
        self.caminho_template = Path(caminho_template)
        if not self.caminho_template.exists():
            raise FileNotFoundError(f'Template não encontrado: {caminho_template}')

        self.doc = Document(str(self.caminho_template))
        self.tabela_evidencias = None
        self._encontrar_tabela_evidencias()

    def _encontrar_tabela_evidencias(self):
        """
        Localiza a tabela de evidências.
        A tabela de evidências é a maior tabela do documento
        (78 linhas × 2 colunas no template base).
        """
        maior = None
        maior_tamanho = 0
        for t in self.doc.tables:
            tamanho = len(t.rows) * len(t.columns)
            if tamanho > maior_tamanho:
                maior_tamanho = tamanho
                maior = t
        self.tabela_evidencias = maior
        logger.info(f'Tabela de evidências encontrada: {len(maior.rows)} linhas × {len(maior.columns)} colunas')

    def substituir_id_projeto(self, projeto_id: str):
        """
        Substitui o placeholder 'Projeto: XXXXXXX' pelo ID do projeto.
        Substitui apenas o padrão 'Projeto: XXXXXXX' para não afetar
        outros campos como 'Município: XXXXXXX'.
        """
        for paragrafo in self.doc.paragraphs:
            texto_completo = paragrafo.text
            if 'Projeto:' in texto_completo and 'XXXXXXX' in texto_completo:
                for run in paragrafo.runs:
                    if 'XXXXXXX' in run.text:
                        run.text = run.text.replace('XXXXXXX', projeto_id)
                        logger.info(f'Placeholder substituído: XXXXXXX → {projeto_id}')

    def _obter_dimensoes_celula(self, cell) -> tuple:
        """
        Retorna (largura, altura) da célula em EMUs para dimensionar a imagem.
        """
        # Largura da célula
        largura_tc = cell.width if cell.width else Inches(2.5)
        largura_emu = largura_tc

        # Altura: usa a altura da linha
        row = cell._element.getparent()
        if row is not None:
            trPr = row.find(qn('w:trPr'))
            if trPr is not None:
                trHeight = trPr.find(qn('w:trHeight'))
                if trHeight is not None:
                    altura_val = int(trHeight.get(qn('w:val')))
                    altura_emu = Emu(altura_val * 914400 // 1440)  # dxa para EMU
                    return (largura_emu, altura_emu)

        return (largura_emu, Emu(largura_emu * 3 // 4))

    def _inserir_imagem_celula(self, cell, caminho_imagem: str):
        """
        Insere uma imagem em uma célula da tabela, ajustando para caber no espaço
        sem deformar a proporção.
        """
        try:
            # Abre a imagem para obter dimensões originais
            with PILImage.open(caminho_imagem) as img:
                largura_orig, altura_orig = img.size

            img_ratio = largura_orig / altura_orig

            # Obtém dimensões disponíveis na célula
            cell_largura_emu, cell_altura_emu = self._obter_dimensoes_celula(cell)

            cell_largura_pt = cell_largura_emu / 914400 * 72
            cell_altura_pt = cell_altura_emu / 914400 * 72

            # Reduz 10% para margem interna
            max_largura_pt = cell_largura_pt * 0.85
            max_altura_pt = cell_altura_pt * 0.85

            # Calcula dimensões mantendo proporção
            if img_ratio > (max_largura_pt / max_altura_pt):
                largura_final_pt = max_largura_pt
                altura_final_pt = max_largura_pt / img_ratio
            else:
                altura_final_pt = max_altura_pt
                largura_final_pt = max_altura_pt * img_ratio

            # Limpa o parágrafo e insere a imagem
            paragraph = cell.paragraphs[0]
            # Remove runs existentes
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)

            run = paragraph.add_run()
            run.add_picture(caminho_imagem, width=int(largura_final_pt * 914400 / 72))

            logger.debug(f'Imagem inserida: {Path(caminho_imagem).name} ({largura_final_pt:.0f}×{altura_final_pt:.0f} pt)')

        except Exception as e:
            logger.error(f'Erro ao inserir imagem {caminho_imagem} na célula: {e}')
            paragraph = cell.paragraphs[0]
            paragraph.text = f'[Erro: {Path(caminho_imagem).name}]'

    def _contar_linhas_foto(self) -> int:
        """
        Conta quantas linhas de foto (altura >= 6520 dxa) existem na tabela.
        """
        tbl = self.tabela_evidencias._tbl
        linhas = tbl.findall(qn('w:tr'))
        count = 0
        for row in linhas:
            trPr = row.find(qn('w:trPr'))
            if trPr is not None:
                trHeight = trPr.find(qn('w:trHeight'))
                if trHeight is not None:
                    h_val = int(trHeight.get(qn('w:val')))
                    if h_val >= 6520:
                        count += 1
        return count

    def _clonar_linhas_pagina(self, rows: list, num_paginas_adicionais: int):
        """
        Clona o padrão de linhas de uma página de evidências.
        O padrão são 4 linhas consecutivas que formam uma página
        com 2 linhas de foto e 2 espaçadores.
        """
        if not rows or num_paginas_adicionais <= 0:
            return

        # Encontra a última página (últimas 4 linhas que formam uma página)
        # Procura pelo padrão: [foto, spacer, foto, spacer]
        # A última linha de foto está na penúltima posição entre as foto-rows
        photo_indices = [i for i, r in enumerate(rows) if self._is_foto_row(r)]
        if len(photo_indices) < 2:
            logger.warning('Padrão de página não encontrado para clonagem')
            return

        # Usa o último par de linhas de foto como referência
        last_photo_idx = photo_indices[-1]
        second_last_photo_idx = photo_indices[-2]

        # As 4 linhas da página: [second_last_photo, spacer, last_photo, spacer_next]
        template_indices = [
            second_last_photo_idx,
            second_last_photo_idx + 1,
            last_photo_idx,
            last_photo_idx + 1
        ]

        # Garante que os índices estão dentro do range
        template_indices = [i for i in template_indices if i < len(rows)]

        template_rows = [rows[i] for i in template_indices]
        parent = template_rows[0].getparent()

        for _ in range(num_paginas_adicionais):
            for tr in template_rows:
                parent.append(deepcopy(tr))

        logger.info(f'{num_paginas_adicionais} página(s) de evidências duplicada(s)')

    def _is_foto_row(self, row) -> bool:
        """Verifica se uma linha é destinada a fotos (altura >= 6520 dxa)."""
        trPr = row.find(qn('w:trPr'))
        if trPr is not None:
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is not None:
                h_val = int(trHeight.get(qn('w:val')))
                return h_val >= 6520
        return False

    def _determinar_grupos_pagina(self) -> list:
        """
        Determina os grupos de linhas que formam cada página de evidências.
        Retorna uma lista de listas, onde cada sublista contém os índices
        das linhas que formam uma página.

        O padrão do template é:
        - Foto: altura 6520 (dxa)
        - Espaçador: altura 56 ou 283 (dxa)
        - Foto: altura 6520 (dxa)
        - Espaçador: altura 56 (dxa)
        """
        tbl = self.tabela_evidencias._tbl
        todas_linhas = tbl.findall(qn('w:tr'))

        # Encontra linhas com altura de foto (6520 dxa)
        linhas_foto = []
        for i, row in enumerate(todas_linhas):
            trPr = row.find(qn('w:trPr'))
            if trPr is not None:
                trHeight = trPr.find(qn('w:trHeight'))
                if trHeight is not None:
                    h_val = int(trHeight.get(qn('w:val')))
                    if h_val >= 6520:
                        linhas_foto.append(i)

        # Agrupa em páginas de 2 linhas de foto cada
        paginas = []
        for i in range(0, len(linhas_foto), PHOTO_ROWS_PER_PAGE):
            grupo = linhas_foto[i:i + PHOTO_ROWS_PER_PAGE]
            if grupo:
                # Inclui também os espaçadores entre as fotos
                primeiro = grupo[0]
                ultimo = grupo[-1]
                # Encontra o espaçador após a última foto
                pagina_indices = list(range(primeiro, ultimo + 2))  # +1 para incluir espaçador após
                paginas.append(pagina_indices)

        return paginas

    def preencher_evidencias(self, fotos: list[ph.PhotoInfo], callback_status=None):
        """
        Preenche a tabela de evidências com as fotos.
        Distribui as fotos em 2 colunas × 2 linhas por página.

        Args:
            fotos: Lista de PhotoInfo ordenadas cronologicamente
            callback_status: Função opcional para callback de progresso
        """
        total_fotos = len(fotos)
        logger.info(f'Preenchendo {total_fotos} fotos na tabela de evidências')

        tbl = self.tabela_evidencias

        # Calcula a capacidade real baseada nas linhas de foto
        photo_row_count = self._contar_linhas_foto()
        capacidade_atual = photo_row_count * PHOTOS_PER_ROW

        logger.info(f'Capacidade atual: {capacidade_atual} fotos ({photo_row_count} linhas de foto)')

        # Se precisar de mais fotos, duplica páginas
        if total_fotos > capacidade_atual:
            fotos_extras = total_fotos - capacidade_atual
            paginas_extras = (fotos_extras + PHOTOS_PER_PAGE - 1) // PHOTOS_PER_PAGE
            logger.info(f'Adicionando {paginas_extras} página(s) extras para {fotos_extras} foto(s)')

            todas_linhas = tbl._tbl.findall(qn('w:tr'))
            self._clonar_linhas_pagina(todas_linhas, paginas_extras)

        # Re-obter as linhas após possível expansão
        todas_linhas = tbl._tbl.findall(qn('w:tr'))

        # Mapeia células para fotos
        foto_idx = 0
        linha_idx = 0

        while foto_idx < total_fotos and linha_idx < len(todas_linhas):
            row = todas_linhas[linha_idx]

            if self._is_foto_row(row):
                cells = row.findall(qn('w:tc'))
                for ci in range(len(cells)):
                    if foto_idx >= total_fotos:
                        break

                    cell = tbl.cell(linha_idx, ci)
                    foto = fotos[foto_idx]
                    self._inserir_imagem_celula(cell, str(foto.caminho))

                    logger.info(f'Foto {foto_idx + 1}/{total_fotos}: {foto.nome} → linha {linha_idx}, coluna {ci}')

                    if callback_status:
                        callback_status(foto_idx + 1, total_fotos, foto.nome)

                    foto_idx += 1

                if foto_idx >= total_fotos:
                    break

            linha_idx += 1

        logger.info(f'Preenchimento concluído: {foto_idx} fotos inseridas')

    def salvar(self, caminho_saida: str) -> str:
        """
        Salva o documento preenchido no caminho especificado.
        Retorna o caminho absoluto do arquivo salvo.
        """
        caminho = Path(caminho_saida)
        caminho.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(caminho))
        logger.info(f'Documento salvo: {caminho}')
        return str(caminho.absolute())
